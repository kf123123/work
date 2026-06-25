# -*- coding: utf-8 -*-
"""
中期报告生成脚本 — 车辆行驶远程控制系统设计
生成 Word 文档，参考开题报告模板格式
"""

import os
from docx import Document
from docx.shared import Pt, Cm, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ============ 路径配置 ============
OUTPUT_DIR = "D:/UserData/Desktop/practice/report"
RESOURCE_DIR = "D:/UserData/Desktop/practice/resource"
OUTPUT_FILE = os.path.join(OUTPUT_DIR, "车辆行驶远程控制系统_中期报告_v2.docx")

# ============ 学生信息 ============
STUDENTS = [
    ("2306024123", "李兆新"),
    ("2306024131", "李科锋"),
    ("2306024132", "白宇"),
]

TEAM_NO = "第三组"
INSTRUCTOR = "王红亮"
REPORT_DATE = "2026年6月"

# ============ 辅助函数 ============

def set_cell_shading(cell, color):
    """设置单元格底色"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_run_font(run, font_name_cn="宋体", font_name_en="Times New Roman", size=12, bold=False):
    """设置 run 字体"""
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = font_name_en
    rpr = run._element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rpr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name_cn)
    rFonts.set(qn('w:ascii'), font_name_en)
    rFonts.set(qn('w:hAnsi'), font_name_en)

def set_paragraph_spacing(paragraph, line_spacing=1.5, space_before=0, space_after=0):
    """设置段落间距"""
    pf = paragraph.paragraph_format
    pf.line_spacing = line_spacing
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)

def add_heading1(doc, text):
    """添加一级标题（黑体 小三/15pt 加粗）"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(p, line_spacing=1.5, space_before=12, space_after=6)
    run = p.add_run(text)
    set_run_font(run, "黑体", "Times New Roman", size=15, bold=True)
    return p

def add_heading2(doc, text):
    """添加二级标题（黑体 四号/14pt 加粗）"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(p, line_spacing=1.5, space_before=8, space_after=4)
    run = p.add_run(text)
    set_run_font(run, "黑体", "Times New Roman", size=14, bold=True)
    return p

def add_heading3(doc, text):
    """添加三级标题（黑体 小四/12pt 加粗）"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(p, line_spacing=1.5, space_before=6, space_after=3)
    run = p.add_run(text)
    set_run_font(run, "黑体", "Times New Roman", size=12, bold=True)
    return p

def add_body(doc, text, indent_first=True, font_size=12):
    """添加正文段落（宋体 小四/12pt）"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_spacing(p, line_spacing=1.5)
    if indent_first:
        p.paragraph_format.first_line_indent = Pt(24)  # 两个字符缩进
    run = p.add_run(text)
    set_run_font(run, "宋体", "Times New Roman", size=font_size)
    return p

def add_body_with_bold_prefix(doc, prefix, text, font_size=12):
    """添加正文，开头部分加粗"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_spacing(p, line_spacing=1.5)
    p.paragraph_format.first_line_indent = Pt(24)
    run1 = p.add_run(prefix)
    set_run_font(run1, "宋体", "Times New Roman", size=font_size, bold=True)
    run2 = p.add_run(text)
    set_run_font(run2, "宋体", "Times New Roman", size=font_size)
    return p

def add_figure_caption(doc, text):
    """添加图题（居中，加粗）"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, line_spacing=1.2, space_before=4, space_after=8)
    run = p.add_run(text)
    set_run_font(run, "宋体", "Times New Roman", size=10, bold=True)
    return p

def add_code_block(doc, code_text, font_size=8):
    """添加代码块（等宽字体，灰色背景用文字模拟）"""
    lines = code_text.strip().split('\n')
    for line in lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_paragraph_spacing(p, line_spacing=1.0, space_before=0, space_after=0)
        # 设置左右缩进
        p.paragraph_format.left_indent = Cm(1)
        run = p.add_run(line)
        set_run_font(run, "Courier New", "Courier New", size=font_size, bold=False)

def add_image_with_caption(doc, img_path, caption, width_cm=14):
    """插入图片并添加图题"""
    if os.path.exists(img_path):
        try:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(img_path, width=Cm(width_cm))
        except Exception as e:
            add_body(doc, f"[图片插入失败: {img_path}]")
    else:
        add_body(doc, f"[图片文件未找到: {img_path}]")
    add_figure_caption(doc, caption)

def add_table_with_data(doc, headers, rows, col_widths=None):
    """添加表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    for ci, h in enumerate(headers):
        cell = table.rows[0].cells[ci]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_run_font(run, "黑体", "Times New Roman", size=10, bold=True)

    # 数据行
    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            set_run_font(run, "宋体", "Times New Roman", size=10)

    return table


# ============ 生成文档 ============

def generate_report():
    doc = Document()

    # ---- 页面设置 ----
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    # ======================== 封面 ========================
    # 空行
    for _ in range(4):
        p = doc.add_paragraph()
        set_paragraph_spacing(p, line_spacing=1.5)

    # 主标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("电子科学与技术专业")
    set_run_font(run, "黑体", "Times New Roman", size=26, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, line_spacing=1.5, space_before=8)
    run = p.add_run("系统方案综合实践与设计")
    set_run_font(run, "黑体", "Times New Roman", size=26, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, line_spacing=1.5, space_before=8)
    run = p.add_run("中期报告")
    set_run_font(run, "黑体", "Times New Roman", size=28, bold=True)

    # 空行
    for _ in range(2):
        p = doc.add_paragraph()
        set_paragraph_spacing(p, line_spacing=1.5)

    # 项目名称
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, line_spacing=1.8)
    run = p.add_run("设计题目：车辆行驶远程控制系统设计")
    set_run_font(run, "黑体", "Times New Roman", size=18, bold=False)

    # 空行
    for _ in range(2):
        doc.add_paragraph()

    # 学生信息表格
    table = doc.add_table(rows=3 + 3, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    info_data = [
        ["", "", "", ""],
        ["组员姓名", "李兆新", "学    号", "2306024123"],
        ["组员姓名", "李科锋", "学    号", "2306024131"],
        ["组员姓名", "白  宇", "学    号", "2306024132"],
        ["组    别", TEAM_NO, "", ""],
        ["指导教师", INSTRUCTOR, "日    期", REPORT_DATE],
    ]

    for ri, row_data in enumerate(info_data):
        for ci, val in enumerate(row_data):
            cell = table.rows[ri].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(val)
            set_run_font(run, "宋体", "Times New Roman", size=12)

    # 合并空行
    table.rows[0].cells[0].merge(table.rows[0].cells[3])

    # 分页
    doc.add_page_break()

    # ======================== 正文 ========================

    # ===== 一、项目概述 =====
    add_heading1(doc, "一、项目概述")

    add_heading2(doc, "1.1 项目背景")
    add_body(doc, "随着物联网技术和无线通信技术的快速发展，远程控制系统在智能交通、物流配送、环境监测等领域得到了广泛应用。本项目旨在设计并实现一套车辆行驶远程控制系统，通过无线网络实现上位机对车辆的运动控制和实时视频监控，同时具备定位追踪功能。")

    add_heading2(doc, "1.2 任务书要求")
    add_body(doc, "根据任务书要求，本系统需要完成以下设计内容：")
    add_body(doc, "（1）上位机监控界面设计：使用 PyQt5 开发图形化用户界面，包含实时视频显示、车辆运动控制面板、状态信息显示和定位地图等功能。")
    add_body(doc, "（2）视频传输功能：通过 Wi-Fi 网络实现下位机摄像头画面到上位机的实时传输，保证画面流畅、延迟低。")
    add_body(doc, "（3）车辆运动控制：上位机可通过按钮或键盘控制车辆执行前进、后退、左转、右转等多方向运动，并支持速度调节。")
    add_body(doc, "（4）定位功能：通过 UWB 定位模块获取车辆位置信息，在上位机地图上实时显示车辆位置和运动轨迹。")

    add_heading2(doc, "1.3 项目组成员与分工")
    add_body(doc, "本项目由三名同学共同完成，具体分工如下：")

    add_table_with_data(doc,
        ["学号", "姓名", "主要职责"],
        [
            ["2306024123", "李兆新", "UWB定位模块数据解析、文档撰写与报告整理"],
            ["2306024131", "李科锋", "视频流传输、上位机主程序开发、与下位机通信协议对接"],
            ["2306024132", "白  宇", "下位机电机驱动与运动控制、下位机协议对接与联调"],
        ]
    )

    # ===== 二、系统总体方案设计 =====
    add_heading1(doc, "二、系统总体方案设计")

    add_heading2(doc, "2.1 系统架构")
    add_body(doc, '本系统采用"上位机 + 下位机"的分布式架构。上位机为一台运行 Windows 系统的 PC，负责提供图形化监控界面；下位机包括 K230 开发板（摄像头视频采集与推流）和 ESP32 开发板（电机驱动控制与定位数据采集）。上位机与下位机之间通过 Wi-Fi 无线网络进行通信。')

    add_body(doc, "系统整体架构如下：")
    add_body(doc, "上位机（PC/PyQt5）⟷ Wi-Fi ⟷ 下位机（K230 + ESP32）")
    add_body(doc, "其中，视频数据采用 UDP 协议传输（端口 8080），控制指令和定位数据采用 TCP 协议传输（端口 8888）。视频传输方向为 K230 → PC，控制指令方向为 PC → ESP32，定位数据方向为 ESP32 → PC。")

    add_heading2(doc, "2.2 通信方案设计")
    add_body(doc, "系统采用两种通信协议以满足不同数据类型的需求：")

    add_body_with_bold_prefix(doc, "（1）视频传输（UDP协议）：",
        "摄像头采集的画面经过 JPEG 压缩后，按每帧最大 1400 字节进行分片，通过 UDP 协议发送至上位机。上位机接收后根据帧 ID 和分片序号重组数据并解码显示。UDP 协议的无连接特性保证了视频传输的低延迟。")

    add_body_with_bold_prefix(doc, "（2）控制与定位（TCP协议）：",
        "控制指令以 JSON 格式通过 TCP 连接发送至 ESP32，确保指令可靠到达。ESP32 解析指令后驱动电机执行相应动作。定位数据同样以 JSON 格式经 TCP 连接定时推送至上位机，实现车辆位置的实时更新。")

    add_heading2(doc, "2.3 通信协议定义")

    add_body(doc, "本系统定义了三种通信协议，分别用于控制指令传输、定位数据传输和视频流传输。所有通信均基于 Wi-Fi 无线网络。")

    add_heading3(doc, "2.3.1 控制指令协议（PC → ESP32，TCP JSON）")
    add_body(doc, "上位机通过 TCP 连接向 ESP32 发送控制指令，每条指令为一行 JSON 数据，以换行符分隔。指令的 T 字段表示命令类型，不同命令携带不同的参数字段：")

    add_table_with_data(doc,
        ["T值", "命令", "参数字段", "说明"],
        [
            ["11", "PWM控制", "L, R", "直接设置左右电机PWM值（-255~255）"],
            ["115", "停止", "无", "立即停止所有电机"],
            ["116", "前进", "pwm", "以指定PWM值前进（默认150）"],
            ["117", "后退", "pwm", "以指定PWM值后退"],
            ["118", "左转", "pwm", "以指定PWM值原地左转"],
            ["119", "右转", "pwm", "以指定PWM值原地右转"],
            ["200", "查询定位", "无", "请求ESP32立即发送一次定位数据"],
            ["201", "定位推送控制", "en", "开启(en=1)或关闭(en=0)UWB自动推送"],
            ["999", "心跳", "无", "心跳包，维持连接"],
        ]
    )

    add_body(doc, "控制指令JSON格式示例：")
    add_code_block(doc, """// 前进，PWM=150
{"T": 116, "pwm": 150}

// 直接PWM控制：左轮200，右轮-200（后退）
{"T": 11, "L": 200, "R": -200}

// 停止
{"T": 115}

// 开启UWB自动推送
{"T": 201, "en": 1}""")

    add_heading3(doc, "2.3.2 定位数据协议（ESP32 → PC，TCP JSON）")
    add_body(doc, "ESP32 定时（默认 5Hz = 200ms 间隔）向上位机推送定位数据，或在上位机请求时单次发送。数据同样以一行 JSON 格式传输：")

    add_table_with_data(doc,
        ["字段", "类型", "说明"],
        [
            ["T", "int", "消息类型，固定为200"],
            ["x", "float", "车辆X坐标（单位：米）"],
            ["y", "float", "车辆Y坐标（单位：米）"],
            ["valid", "bool", "定位数据是否有效"],
            ["t", "int", "时间戳（毫秒）"],
        ]
    )

    add_body(doc, "定位数据JSON格式示例：")
    add_code_block(doc, """{"T": 200, "x": 1.234, "y": 0.567, "valid": true, "t": 12345}""")
    add_body(doc, "上位机 TcpReceiveThread 接收到该数据后，解析坐标并在定位地图上更新车辆位置。valid 字段用于标识定位数据是否超时（超过 1 秒未收到 UWB 数据则视为无效）。")

    add_heading3(doc, "2.3.3 视频传输协议（K230 → PC，UDP分片）")
    add_body(doc, "K230 下位机将摄像头采集的画面经 JPEG 压缩后，通过 UDP 协议分片发送至上位机。由于单次 UDP 数据包大小限制（以太网 MTU 约 1500 字节），每帧 JPEG 数据按 1400 字节分片传输，每个分片携带 6 字节包头：")

    add_table_with_data(doc,
        ["字段", "长度", "类型", "说明"],
        [
            ["frame_id", "2字节", "unsigned short (大端)", "帧序号，递增循环"],
            ["total_chunks", "2字节", "unsigned short (大端)", "当前帧总分片数"],
            ["chunk_id", "2字节", "unsigned short (大端)", "当前分片序号（从0开始）"],
            ["payload", "N字节", "JPEG数据", "JPEG图像数据片段"],
        ]
    )

    add_body(doc, "UDP数据包结构示意：")
    add_code_block(doc, """// UDP数据包结构（每包 ≤ 1406 字节）
// [frame_id:2字节] [total_chunks:2字节] [chunk_id:2字节] [JPEG payload:N字节]
//
// 例如一帧JPEG数据为4200字节，则分为3片发送：
//   包1: frame_id=1, total_chunks=3, chunk_id=0, payload=[0~1399]
//   包2: frame_id=1, total_chunks=3, chunk_id=1, payload=[1400~2799]
//   包3: frame_id=1, total_chunks=3, chunk_id=2, payload=[2800~4199]""")

    add_body(doc, "上位机 VideoReceiveThread 收到分片后，根据 frame_id 和 total_chunks 分配缓冲区，收集齐所有分片后拼接为完整 JPEG 数据，调用 OpenCV 解码并在视频窗口中显示。")

    add_heading2(doc, "2.4 硬件选型")
    add_table_with_data(doc,
        ["组件", "型号/方案", "用途"],
        [
            ["主控MCU", "ESP32 Dev Module", "电机驱动控制、WiFi通信、UWB数据采集"],
            ["摄像头模块", "K230（内置ISP）", "视频图像采集与JPEG压缩编码"],
            ["电机驱动", "L298N / TB6612", "驱动直流电机"],
            ["定位模块", "Nooploop LinkTrack", "UWB定位，获取车辆坐标"],
            ["底盘", "4轮差速小车底盘", "承载各模块"],
        ]
    )

    # ===== 三、硬件设计与实现 =====
    add_heading1(doc, "三、硬件设计与实现")

    add_heading2(doc, "3.1 电路接口设计")
    add_body(doc, "ESP32 主控板与各模块之间的引脚连接关系如下：")

    add_body_with_bold_prefix(doc, "电机驱动接口（左右两路直流电机）：", "")
    add_table_with_data(doc,
        ["信号", "ESP32引脚", "功能说明"],
        [
            ["PIN_PWMA", "GPIO25", "左电机 PWM 调速"],
            ["PIN_AIN2", "GPIO17", "左电机方向控制 A"],
            ["PIN_AIN1", "GPIO21", "左电机方向控制 B"],
            ["PIN_BIN1", "GPIO22", "右电机方向控制 A"],
            ["PIN_BIN2", "GPIO23", "右电机方向控制 B"],
            ["PIN_PWMB", "GPIO26", "右电机 PWM 调速"],
        ]
    )

    add_body(doc, "PWM 频率配置为 100kHz，分辨率为 8 位（0~255）。通过两个 H 桥电路分别控制左右电机的正反转和转速，实现差速驱动。当左右轮转速相同时实现前进/后退，转速不同时实现转向，转速相反时实现原地旋转。")

    add_body_with_bold_prefix(doc, "UWB 定位模块接口：", "")
    add_table_with_data(doc,
        ["信号", "ESP32引脚", "说明"],
        [
            ["UART2 RX", "GPIO16", "UWB 数据接收（单向）"],
            ["UART2 TX", "未使用", "GPIO17 已被电机占用"],
            ["波特率", "921600", "高速串口通信"],
        ]
    )

    add_body(doc, "UWB 模块通过 UART2 与 ESP32 通信，使用 921600 波特率的高速串口模式。由于 GPIO17 已被电机方向控制引脚占用，故仅使用单向接收模式接收 UWB 数据。")

    add_heading2(doc, "3.2 电机驱动原理")
    add_body(doc, "电机采用差速驱动方式，通过独立控制左右两个电机的转速和方向来实现车辆的各种运动。驱动逻辑如下：")

    add_table_with_data(doc,
        ["运动状态", "左电机", "右电机", "说明"],
        [
            ["前进", "正转（PWM）", "正转（PWM）", "两轮同速正转"],
            ["后退", "反转（PWM）", "反转（PWM）", "两轮同速反转"],
            ["原地左转", "反转（PWM）", "正转（PWM）", "两轮反向等速旋转"],
            ["原地右转", "正转（PWM）", "反转（PWM）", "两轮反向等速旋转"],
            ["左前转", "停止", "正转（PWM）", "右轮单独正转"],
            ["右前转", "正转（PWM）", "停止", "左轮单独正转"],
            ["停止", "停止", "停止", "全部制动"],
        ]
    )

    # ===== 四、软件设计与实现 =====
    add_heading1(doc, "四、软件设计与实现")

    add_heading2(doc, "4.1 上位机软件设计")
    add_body(doc, "上位机基于 PyQt5 框架开发，采用事件驱动模型。主程序由三个核心模块组成：")

    add_heading3(doc, "4.1.1 视频接收线程（VideoReceiveThread）")
    add_body(doc, "该线程通过 UDP 协议监听 8080 端口，接收来自 K230 的 JPEG 分片数据。每帧数据包含帧 ID、总分片数和分片序号，接收完一帧的所有分片后重组为完整的 JPEG 图像，解码并转换为 QImage 后通过信号机制发送到主界面显示。同时计算并更新视频帧率。")

    add_body(doc, "关键代码与注释：")
    add_code_block(doc, """class VideoReceiveThread(QThread):
    def __init__(self):
        super().__init__()
        self.udp_socket = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
        self.udp_socket.bind(('0.0.0.0', 8080))   # 监听8080 UDP端口
        self.frame_buffer = {}                      # 分片缓冲区

    def run(self):
        while self.running:
            data, addr = self.udp_socket.recvfrom(2048)  # 接收UDP数据包
            # 解析6字节包头: frame_id(2) + total_chunks(2) + chunk_id(2)
            frame_id, total_chunks, chunk_id = struct.unpack("!HHH", data[:6])
            payload = data[6:]                   # 剩余为JPEG数据片

            # 存入缓冲区
            if frame_id not in self.frame_buffer:
                self.frame_buffer[frame_id] = [None] * total_chunks
            self.frame_buffer[frame_id][chunk_id] = payload

            # 分片收齐后重组
            if all(p is not None for p in self.frame_buffer[frame_id]):
                jpeg_data = b"".join(self.frame_buffer[frame_id])
                nparr = np.frombuffer(jpeg_data, np.uint8)
                frame = cv2.imdecode(nparr, cv2.IMREAD_COLOR)
                # 转换并发射信号到主界面显示""")

    add_heading3(doc, "4.1.2 定位数据接收线程（TcpReceiveThread）")
    add_body(doc, "该线程负责通过 TCP 连接接收 ESP32 发送的 UWB 定位数据。数据以 JSON 格式逐行传输，包含车辆坐标（x, y）、有效标志和时间戳等信息。解析后通过信号发送至主界面的地图模块进行坐标更新和轨迹绘制。")

    add_body(doc, "关键代码与注释：")
    add_code_block(doc, """class TcpReceiveThread(QThread):
    def run(self):
        while self.running:
            data = self.tcp_socket.recv(1024)         # 接收TCP数据
            self.buffer += data.decode('utf-8')
            while "\\n" in self.buffer:               # 按行解析JSON
                line, self.buffer = self.buffer.split("\\n", 1)
                data_dict = json.loads(line)
                if data_dict.get("T") == 200:         # T=200 为定位数据
                    x, y = data_dict.get('x'), data_dict.get('y')
                    self.location_received_signal.emit(data_dict)""")

    add_heading3(doc, "4.1.3 主控窗口（ControlWindow）")
    add_body(doc, "主界面整合了视频显示、8方向控制面板、速度滑块、状态表格和定位地图等多个功能模块。支持鼠标点击按钮和键盘两种控制方式。键盘采用 WASD 经典布局（W=前进，S=后退，A=左转，D=右转，空格=停止），同时支持 Q/E/Z/C 实现四角方向。速度滑块可实时调节 PWM 占空比（0~200），运动中调整速度立即生效。")

    add_body(doc, "运动控制指令采用 JSON 格式发送，使用 T 字段标识命令类型，L/R 字段分别指定左右电机 PWM 值：")
    add_code_block(doc, """# 运动控制指令编码（T值）
CMD = {
    "STOP": 115, "FORWARD": 116, "BACKWARD": 117,
    "LEFT": 118, "RIGHT": 119,
    "FWD_LEFT": 120, "FWD_RIGHT": 121,
    "BWD_LEFT": 122, "BWD_RIGHT": 123,
}

# 差速驱动映射：每个方向对应的(L_ratio, R_ratio)
DIR_MOTOR = {
    CMD["FORWARD"]:   (1.0, 1.0),     # 前进
    CMD["BACKWARD"]:  (-1.0, -1.0),   # 后退
    CMD["LEFT"]:      (-1.0, 1.0),    # 原地左转
    CMD["RIGHT"]:     (1.0, -1.0),    # 原地右转
    CMD["STOP"]:      (0, 0),         # 停止
}

def _make_motor_cmd(self, direction):
    # 根据方向计算左右轮PWM值
    lr, rr = DIR_MOTOR[direction]
    l = max(-255, min(255, int(lr * self.current_pwm)))
    r = max(-255, min(255, int(rr * self.current_pwm)))
    return {"T": 11, "L": l, "R": r}""")

    add_heading2(doc, "4.2 上位机界面布局")
    add_body(doc, "上位机界面采用左右分栏布局，左侧为主操作区（视频画面 + 控制面板），右侧为信息显示区（车辆状态 + 定位地图），整体采用暗色主题风格，具有较好的视觉效果。")

    # 插入上位机界面截图
    add_image_with_caption(doc,
        os.path.join(RESOURCE_DIR, "pyqt.jpg"),
        "图1 上位机监控主界面", width_cm=14)

    add_heading2(doc, "4.3 下位机软件设计（ESP32固件）")
    add_body(doc, "ESP32 下位机固件使用 Arduino 框架开发，集成电机驱动、WiFi/TCP 通信和 UWB 定位数据解析三大功能模块。主程序在 loop() 中循环执行三个任务：")

    add_body(doc, "（1）读取 UWB 串口数据并解析定位坐标；")
    add_body(doc, "（2）处理 TCP 客户端连接请求、接收控制指令、定时推送定位数据；")
    add_body(doc, "（3）处理串口调试命令。")

    add_body(doc, "主程序流程：")
    add_code_block(doc, """void setup() {
    chassisInit();           // 初始化电机引脚和PWM
    initUwbSerial();         // 初始化UWB串口(UART2)
    wifiTcpInit();           // 连接WiFi并启动TCP Server
}

void loop() {
    uwbSerialRead();         // 1. 高频读取UWB数据
    wifiTcpLoop();           // 2. 处理TCP连接、命令、UWB推送
    handleSerial();          // 3. 处理串口调试命令
}""")

    add_heading2(doc, "4.4 电机驱动控制逻辑")
    add_body(doc, "电机驱动模块通过 H 桥电路控制电机的正反转和速度。左右两个电机的方向由四个控制引脚的电平组合决定，速度由 PWM 占空比调节。代码实现了对 ESP32 Arduino 3.x 新版 API（ledcAttach/ledcWrite）的适配。")

    add_code_block(doc, """// 左轮驱动
static inline void driveA(int pwm) {
    if (MOTOR_DIR_REVERSED) {
        if (pwm >= 0) { digitalWrite(PIN_AIN1, LOW);  digitalWrite(PIN_AIN2, HIGH); }
        else          { digitalWrite(PIN_AIN1, HIGH); digitalWrite(PIN_AIN2, LOW);  pwm = -pwm; }
    }
    ledcWrite(PIN_PWMA, pwm);           // ESP32新API：直接按引脚写入PWM
}

// 右轮驱动
static inline void driveB(int pwm) {
    if (pwm >= 0) { digitalWrite(PIN_BIN1, HIGH); digitalWrite(PIN_BIN2, LOW); }
    else          { digitalWrite(PIN_BIN1, LOW);  digitalWrite(PIN_BIN2, HIGH); pwm = -pwm; }
    ledcWrite(PIN_PWMB, pwm);           // 按引脚写入PWM
}

// 差速控制：分别设置左右轮PWM
void motorSet(int pwmLeft, int pwmRight) {
    driveA(pwmLeft);     // 左轮
    driveB(pwmRight);    // 右轮
}""")

    add_heading2(doc, "4.5 K230 摄像头推流")
    add_body(doc, "K230 下位机负责采集摄像头画面，经过 JPEG 压缩（quality=35）后，通过 UDP 协议分片发送到上位机。JPEG 质量参数的优化可在保持图像可辨识度的同时显著提高传输帧率。")

    add_code_block(doc, """# K230摄像头采集与UDP推流
sensor = Sensor()
sensor.reset()
sensor.set_framesize(width=640, height=480)
sensor.set_pixformat(Sensor.RGB565)

while True:
    img = sensor.snapshot()
    jpeg_data = img.compress_for_ide(quality=35)  # JPEG压缩

    frame_id = (frame_id + 1) & 0xFFFF
    total_chunks = (len(jpeg_data) + MAX_UDP_SIZE - 1) // MAX_UDP_SIZE

    for chunk_id in range(total_chunks):
        # 每片: 6字节包头(帧ID+总分片数+分片序号) + JPEG数据
        header = struct.pack("!HHH", frame_id, total_chunks, chunk_id)
        udp_socket.sendto(header + payload, dest_addr)""")

    # ===== 五、定位模块设计与实现 =====
    doc.add_page_break()
    add_heading1(doc, "五、定位模块设计与实现")

    add_heading2(doc, "5.1 UWB 定位原理")
    add_body(doc, "本系统采用 Nooploop LinkTrack 系列 UWB 定位模块，基于到达时间差（TDOA）定位原理。通过在空间中部署已知位置的基站，测量标签到各基站的信号到达时间差，进而解算出标签的位置坐标。系统定位精度可达厘米级（视具体环境而定，通常为 10~30cm）。")

    add_heading2(doc, "5.2 通信协议解析")
    add_body(doc, "UWB 模块通过串口输出 NLink 格式的二进制数据帧。ESP32 在 UART2 上接收数据后，按照 NLink 协议进行帧同步和解析。支持两种数据帧格式：")

    add_body(doc, "（1）Tag_Frame0（功能标记 0x01）：短帧，10 字节，包含位置坐标信息。")
    add_body(doc, "（2）Node_Frame2（功能标记 0x04）：长帧，包含更丰富的节点信息，帧长度由第 2~3 字节指示。")

    add_body(doc, "定位坐标采用 24 位有符号整数编码（int24），单位为毫米，解析时需除以 1000 转换为米。数据帧末尾的校验和用于验证数据完整性。")

    add_code_block(doc, """// NLink协议帧解析
// 帧同步头: 0x55
if (!uwbSynced) {
    if (c == 0x55) { uwbBuf[0] = c; uwbIdx = 1; uwbSynced = true; }
    continue;
}

// Tag_Frame0类型: 10字节定长帧
if (uwbBuf[1] == 0x01 && uwbIdx >= 10) {
    uwbPosX = parseInt24(uwbBuf + 4) / 1000.0f;  // X坐标(mm→m)
    uwbPosY = parseInt24(uwbBuf + 7) / 1000.0f;  // Y坐标(mm→m)
    uwbDataValid = true;
}

// Node_Frame2类型: 可变长帧，长度在字节2-3
if (uwbBuf[1] == 0x04 && uwbIdx >= 4) {
    uwbFrameLen = (uint16_t)uwbBuf[2] | ((uint16_t)uwbBuf[3] << 8);
    if (uwbIdx >= uwbFrameLen && nlinkChecksumOk(uwbBuf, uwbFrameLen)) {
        uwbPosX = parseInt24(uwbBuf + 12) / 1000.0f;
        uwbPosY = parseInt24(uwbBuf + 15) / 1000.0f;
    }
}""")

    add_heading2(doc, "5.3 定位数据显示")
    add_body(doc, "ESP32 解析得到车辆位置后，以 5Hz（200ms 间隔）的频率通过 TCP 连接向上位机推送定位数据。数据格式为 JSON：")
    add_code_block(doc, """{"T": 200, "x": 1.234, "y": 0.567, "valid": true, "t": 12345}""")
    add_body(doc, "上位机接收到定位数据后，在右侧定位地图上以红色圆点标记车辆位置，并以橙色轨迹线绘制历史运动路径，同时状态栏实时显示当前坐标值。")

    add_heading2(doc, "5.4 定位模块现状说明")
    add_body(doc, "在项目实施过程中，UWB 定位模块的通信协议解析代码已完成编写并通过测试，ESP32 能够正确接收并解析 NLink 格式的定位数据，定位数据可通过 TCP 连接正常推送至上位机并在地图上显示。")
    add_body(doc, "然而，在后续的系统联调阶段中，UWB 定位模块因硬件故障（模块损坏）无法继续正常工作。目前正在评估模块更换方案，计划采购同型号模块替换后完成定位功能的最终联调测试。")

    # ===== 六、实验与测试 =====
    add_heading1(doc, "六、实验与测试")

    add_heading2(doc, "6.1 测试原理")
    add_body(doc, "系统测试分为单元测试和联调测试两个阶段。单元测试阶段对各功能模块分别进行验证，包括：上位机界面响应测试、视频传输功能测试、电机驱动控制测试、TCP 通信测试。联调测试阶段将所有模块集成后进行整体功能验证。")

    add_heading2(doc, "6.2 测试方法与结果")

    add_heading3(doc, "6.2.1 上位机功能测试")
    add_body(doc, "使用 tcp_listener.py 模拟下位机 TCP 服务端，验证上位机是否能正确发送控制指令。测试表明，所有 8 个方向按钮和键盘快捷键均可正常发送对应 JSON 指令，速度滑块可实时调节 PWM 值。")

    add_heading3(doc, "6.2.2 视频传输测试")
    add_body(doc, "使用 test.py 在本地生成动态测试画面并发送 UDP 数据包，验证上位机视频接收功能。测试表明，视频画面可正常显示，帧率稳定在 25~30 FPS。使用真实 K230 下位机时，在 quality=35 参数下可达到约 20 FPS 的流畅度。")

    add_heading3(doc, "6.2.3 电机控制测试")
    add_body(doc, "通过串口向 ESP32 发送控制指令，验证电机驱动响应。测试了前进、后退、转向、停止等所有 9 种运动状态，车辆响应正确，PWM 调速功能正常。")

    add_heading3(doc, "6.2.4 UWB 定位测试")
    add_body(doc, "使用 simulate_uwb.py 模拟器生成圆形、方形、8 字形三种运动轨迹，验证上位机定位地图的坐标显示和轨迹绘制功能。测试表明，定位红点可正确跟随模拟坐标移动，橙色轨迹线连续平滑。")

    add_heading2(doc, "6.3 测试现场照片")

    # 插入测试照片
    add_image_with_caption(doc,
        os.path.join(RESOURCE_DIR, "practice1.jpg"),
        "图2 系统联调测试现场（1）", width_cm=12)

    add_image_with_caption(doc,
        os.path.join(RESOURCE_DIR, "practice2.jpg"),
        "图3 系统联调测试现场（2）", width_cm=12)

    add_body(doc, "系统联调测试现场照片如图 2、图 3 所示。实际测试中，车辆可通过上位机远程控制进行前进、后退、转向等动作，视频画面传输流畅，基本达到任务书要求。")

    add_heading2(doc, "6.4 演示视频说明")
    add_body(doc, "项目测试过程中录制了以下演示视频：")
    add_body(doc, "（1）上位机演示视频（pyqt_video.mp4）：展示了上位机界面的完整操作流程，包括视频画面显示、8方向控制、速度调节、地图定位等功能。")
    add_body(doc, "（2）小车行走演示视频（car_control.mp4）：记录了车辆在实际场景中受上位机远程控制的运行情况。")

    # ===== 七、完成情况分析 =====
    add_heading1(doc, "七、完成情况分析")

    add_heading2(doc, "7.1 任务书要求完成情况")
    add_body(doc, "对照任务书的各项要求，项目完成情况分析如下：")

    add_table_with_data(doc,
        ["序号", "任务要求", "完成状态", "说明"],
        [
            ["1", "上位机监控界面设计", "✅ 已完成", "基于PyQt5开发，包含视频、控制、状态、地图四个功能区域"],
            ["2", "实时视频传输功能", "✅ 已完成", "K230→PC的UDP分片传输，帧率约20FPS"],
            ["3", "车辆运动控制功能", "✅ 已完成", "支持8方向控制+速度调节，按钮和键盘双操作模式"],
            ["4", "UWB定位功能", "⚠️ 部分完成", "通信协议解析代码已完成并验证，硬件后续损坏"],
            ["5", "系统联调测试", "🔄 基本完成", "各模块单元测试通过，UWB因硬件问题待修复"],
        ]
    )

    add_heading2(doc, "7.2 未完成项分析")
    add_body(doc, "UWB 定位模块的硬件故障是本项目目前存在的主要问题。该模块在开发阶段已完成通信协议解析和上位机数据显示的完整代码编写，并通过模拟器验证了功能的正确性。但因实际硬件模块损坏，无法在实际场景中完成定位功能的验证和演示。后续计划更换硬件模块后完成此部分工作。")

    # ===== 八、存在问题与改进 =====
    add_heading1(doc, "八、存在问题与改进")

    add_heading2(doc, "8.1 存在的问题")
    add_body(doc, "（1）UWB 定位模块硬件损坏：已完成软件开发和测试的 UWB 定位模块在实际联调阶段出现硬件故障，无法继续使用。")
    add_body(doc, "（2）视频传输帧率受限于 Wi-Fi 环境：在 Wi-Fi 信号较弱时，视频传输帧率会下降，影响实时监控体验。")
    add_body(doc, "（3）车辆控制精度有待优化：开环控制下的电机响应存在一定延迟，定位控制精度有待进一步提升。")

    add_heading2(doc, "8.2 改进措施")
    add_body(doc, "（1）采购同型号 UWB 模块进行替换，重新完成定位功能的联调测试。")
    add_body(doc, "（2）优化 JPEG 压缩参数和 UDP 发送策略，在保证图像质量的前提下进一步提高视频帧率。")
    add_body(doc, "（3）考虑引入 PID 控制算法优化电机控制精度，提高车辆运动控制的稳定性。")

    # ===== 九、总结与展望 =====
    add_heading1(doc, "九、总结与展望")

    add_heading2(doc, "9.1 已完成工作总结")
    add_body(doc, "经过项目组成员的共同努力，车辆行驶远程控制系统设计已取得了阶段性成果。目前已完成的工作包括：")
    add_body(doc, "（1）完成了基于 PyQt5 的上位机监控软件开发，实现了视频显示、8 方向控制、速度调节、定位地图等核心功能，界面设计采用了专业的暗色主题风格。")
    add_body(doc, "（2）完成了基于 K230 的下位机摄像头推流程序开发，实现了 UDP 分片传输方案，保证了视频画面的实时性。")
    add_body(doc, "（3）完成了基于 ESP32 的下位机固件开发，包括电机驱动、WiFi/TCP 通信、UWB 定位数据解析等功能。")
    add_body(doc, "（4）完成了系统联调测试，验证了各功能模块的基本可用性。")

    add_heading2(doc, "9.2 后续工作计划")
    add_body(doc, "（1）更换 UWB 定位模块硬件，完成定位功能在实际场景中的验证测试。")
    add_body(doc, "（2）优化系统整体性能，提升视频传输帧率、降低控制延迟。")
    add_body(doc, "（3）完善系统文档，整理实验数据，为结题验收做准备。")

    # ===== 参考文献 =====
    doc.add_page_break()
    add_heading1(doc, "参考文献")
    refs = [
        "[1] 阮一骏. PyQt5 快速开发与实战[M]. 北京: 机械工业出版社, 2022.",
        "[2] 刘火良, 杨森. ESP32 物联网开发实战[M]. 北京: 人民邮电出版社, 2021.",
        "[3] Nooploop. LinkTrack NLink Protocol Specification[Z]. 2023.",
        "[4] Arduino. ESP32 Arduino Core Documentation[EB/OL]. https://docs.espressif.com/projects/arduino-esp32/.",
        "[5] 朱文斌. 基于UWB的室内定位系统设计与实现[J]. 电子技术应用, 2022, 48(3): 85-89.",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        set_paragraph_spacing(p, line_spacing=1.5)
        p.paragraph_format.first_line_indent = Pt(24)
        run = p.add_run(ref)
        set_run_font(run, "宋体", "Times New Roman", size=11)

    # ---- 保存 ----
    doc.save(OUTPUT_FILE)
    print(f"报告已生成: {OUTPUT_FILE}")


if __name__ == "__main__":
    generate_report()
